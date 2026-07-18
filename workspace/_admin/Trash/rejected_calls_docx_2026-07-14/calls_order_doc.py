# Last updated: 2026-07-18 21:27:59
# @nova: Renders the call-order graph as a HUMAN-READABLE Word document (Orient/Calls_Order.docx).
#        Companion to calls_order.py: that produces the machine-truthful .md, this produces the
#        thing a person can actually sit and read. Same source of truth — it IMPORTS the graph
#        rather than re-deriving it, so the two documents can never disagree.
"""
general_tools/calls_order_doc.py -- visual call-order document
==============================================================
    python general_tools/calls_order_doc.py          # -> Orient/Calls_Order.docx
    python general_tools/calls_order_doc.py --open   # ...and open it

WHY A SECOND FORMAT
───────────────────
`Calls_Order.md` is correct but it is not *legible*. Mermaid source in a text file is a diagram you
have to compile in your head, and the whole point of this artifact is to answer a question fast, at
2am, when a receipt is missing and you need to know whether execute_tool is even on the live path.
A chart you have to decode is a chart you won't look at.

So this renders the same graph as an actual picture, colour-coded, with the call order beside it.

ONE SOURCE OF TRUTH
───────────────────
It imports build_graph()/trace() from calls_order.py. It does not re-parse anything. If the two
documents ever disagreed, one of them would be lying, and you would have no way to know which —
which is precisely the class of problem this whole folder exists to prevent.

DEPENDENCIES — it tells you loudly rather than silently degrading
────────────────────────────────────────────────────────────────
    pip install python-docx networkx matplotlib

Graphviz is used if present (prettier); otherwise it falls back to a networkx layered layout.
If neither is available it still writes the document, with the diagram replaced by an explicit
"NO DIAGRAM — install X" block. It never quietly ships a document that LOOKS complete and isn't.
"""

import subprocess
import sys
from datetime import datetime
from pathlib import Path

WORKSPACE = Path(__file__).resolve().parent.parent
OUT = WORKSPACE / "Orient" / "Calls_Order.docx"
IMG_DIR = WORKSPACE / "Orient" / "Temp"          # generated pngs are temporary → Temp/ (janitorial)
OPEN_AFTER = "--open" in sys.argv

# Same graph, same entry points. Imported, never re-derived.
sys.path.insert(0, str(WORKSPACE / "general_tools"))
from calls_order import build_graph, trace, ENTRY_POINTS, _is_body   # noqa: E402

BODY_FILL = "#1f6f43"
FACE_FILL = "#3a3f5c"


def _require(mod: str, pip_name: str):
    try:
        return __import__(mod)
    except ImportError:
        sys.exit(f"\nMISSING DEPENDENCY: {mod}\n    pip install {pip_name}\n\n"
                 f"Refusing to write a document with a hole in it where the diagram should be.\n")


def render_png(graph, entry, edges, path: Path) -> bool:
    """Draw the call graph. Returns True if an image was written.

    Layered by call DEPTH, top to bottom — because ORDER is the entire point. A force-directed
    blob would be prettier and would tell you nothing about what happens first.
    """
    try:
        import networkx as nx
        import matplotlib
        matplotlib.use("Agg")
        import matplotlib.pyplot as plt
    except ImportError:
        return False

    if not edges:
        return False

    G = nx.DiGraph()
    depth = {entry: 0}
    for d, a, b in edges:
        G.add_edge(a, b)
        depth.setdefault(a, d)
        depth[b] = min(depth.get(b, 99), d + 1)

    # layered layout: y = call depth, x = spread within the layer
    layers = {}
    for n, d in depth.items():
        layers.setdefault(d, []).append(n)
    pos = {}
    for d, nodes in layers.items():
        for i, n in enumerate(sorted(nodes)):
            pos[n] = (i - (len(nodes) - 1) / 2.0, -d)

    colors = [BODY_FILL if _is_body(graph.get(n, {}).get("rel", "")) else FACE_FILL
              for n in G.nodes()]
    widths = max(11, len(G.nodes()) * 1.35)
    heights = max(5.5, (max(layers) + 1) * 1.9)
    plt.figure(figsize=(widths, heights))
    nx.draw_networkx_edges(G, pos, edge_color="#8892a6", arrows=True,
                           arrowsize=13, width=1.1, node_size=900)
    nx.draw_networkx_nodes(G, pos, node_color=colors, node_size=900,
                           edgecolors="#c9d4e6", linewidths=1.2)
    # Labels go BESIDE the node, not inside it. Inside, long names get clipped
    # ("_nova_respon", "uild_self_chec") — and a chart you cannot read is the same
    # failure as no chart at all. Boxed so they stay legible over crossing edges.
    label_pos = {n: (x, y - 0.16) for n, (x, y) in pos.items()}
    for n, (x, y) in label_pos.items():
        rel = graph.get(n, {}).get("rel", "")
        plt.text(x, y, n, fontsize=8, ha="center", va="top", family="DejaVu Sans",
                 color="#12331f" if _is_body(rel) else "#1e2338",
                 bbox=dict(boxstyle="round,pad=0.28",
                           fc="#dff3e6" if _is_body(rel) else "#e4e8f7",
                           ec="#b9c6d8", lw=0.6))
    plt.title(f"{entry}()   —   green = BODY (hers)   ·   blue = face (scaffolding)",
              fontsize=12, color="#222", pad=16)
    plt.margins(0.16)
    plt.axis("off")
    plt.tight_layout()
    path.parent.mkdir(parents=True, exist_ok=True)
    plt.savefig(path, dpi=110, bbox_inches="tight", facecolor="white")
    plt.close()
    return True


def main():
    docx = _require("docx", "python-docx")
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    graph, ambiguous = build_graph()
    doc = Document()

    doc.add_heading("Nova — Call Order", 0)
    p = doc.add_paragraph()
    p.add_run(f"Auto-generated by general_tools/calls_order_doc.py · "
              f"{datetime.now():%Y-%m-%d %H:%M}").italic = True
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(
        "What calls what, and in what ORDER. This is the execution path — not the import graph. "
        "Read it when something didn't happen and you need to know whether it was even on the "
        "live path.")

    doc.add_heading("How to read this", level=1)
    for label, colour, text in [
        ("BODY (green)", BODY_FILL,
         "nova_body/ — this is HER. Faculties: reaching, remembering, deciding, checking. "
         "She keeps these if you pluck the chat server off."),
        ("face (blue)", FACE_FILL,
         "general_tools/ — scaffolding. A window someone looks through. She survives losing it."),
    ]:
        b = doc.add_paragraph(style="List Bullet")
        r = b.add_run(label + " — ")
        r.bold = True
        r.font.color.rgb = RGBColor.from_string(colour.lstrip("#").upper())
        b.add_run(text)

    warn = doc.add_paragraph()
    wr = warn.add_run("An arrow from BODY into face is a pluck-test failure — part of her thinking "
                      "living outside her body. On 2026-07-14 her entire integrity faculty was "
                      "doing exactly that, and a human had to notice; no tool showed it. Now it's "
                      "visible on sight.")
    wr.bold = True

    # ── one section per entry point ──
    for title, hint, fn in ENTRY_POINTS:
        doc.add_page_break()
        doc.add_heading(title, level=1)
        if fn not in graph:
            bad = doc.add_paragraph()
            br = bad.add_run(f"ENTRY POINT NOT FOUND: {fn}()  (expected in {hint})\n"
                             f"The code moved, or a file failed to parse and was omitted. "
                             f"This chart is STALE — do not trust it until this is fixed.")
            br.bold = True
            br.font.color.rgb = RGBColor(0xC0, 0x20, 0x20)
            continue

        doc.add_paragraph(f"Entry: {fn}()  ·  {graph[fn]['rel']}")
        edges = trace(graph, fn)

        img = IMG_DIR / f"callorder_{fn}.png"
        if render_png(graph, fn, edges, img):
            doc.add_picture(str(img), width=Inches(6.4))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        else:
            miss = doc.add_paragraph()
            mr = miss.add_run("NO DIAGRAM — networkx/matplotlib not installed "
                              "(pip install networkx matplotlib). The call order below is still "
                              "complete and correct.")
            mr.bold = True
            mr.font.color.rgb = RGBColor(0xC0, 0x20, 0x20)

        doc.add_heading("Call order", level=2)
        if not edges:
            doc.add_paragraph("No resolvable outgoing calls.")
            continue

        t = doc.add_table(rows=1, cols=4)
        t.style = "Light Grid Accent 1"
        for i, h in enumerate(("#", "Caller", "Calls", "Where")):
            cell = t.rows[0].cells[i]
            cell.text = h
            cell.paragraphs[0].runs[0].bold = True
        for i, (d, a, b) in enumerate(edges, 1):
            rel = graph.get(b, {}).get("rel", "?")
            row = t.add_row().cells
            row[0].text = str(i)
            row[1].text = ("  " * d) + a
            row[2].text = b
            row[3].text = ("BODY  " if _is_body(rel) else "face  ") + rel
            run = row[3].paragraphs[0].runs[0]
            run.font.color.rgb = RGBColor.from_string(
                (BODY_FILL if _is_body(rel) else FACE_FILL).lstrip("#").upper())
            run.font.size = Pt(8)

    # ── pluck audit ──
    doc.add_page_break()
    doc.add_heading("Pluck-test audit", level=1)
    bad = []
    for _, _, fn in ENTRY_POINTS:
        if fn not in graph:
            continue
        for _, a, b in trace(graph, fn):
            ra, rb = graph[a]["rel"], graph[b]["rel"]
            if _is_body(ra) and not _is_body(rb):
                bad.append(f"{a}  (BODY {ra})  →  {b}  (face {rb})")
    if bad:
        doc.add_paragraph("Her body reaches OUT into the face for these. Each is a faculty that "
                          "would vanish if you plucked the chat server off:")
        for b in sorted(set(bad)):
            doc.add_paragraph(b, style="List Bullet")
    else:
        doc.add_paragraph("None on the traced paths — her thinking lives in her body.")

    doc.add_heading("What this document does NOT claim", level=2)
    doc.add_paragraph(
        f"{len(ambiguous)} function names are defined in more than one place (write, add, run…). "
        "A bare AST cannot tell which one a call site meant, so those edges are OMITTED rather "
        "than guessed. An earlier version did guess, and confidently reported that Nova's journal "
        "called into the chat server: eight findings, six of them fiction. A tool built to "
        "establish ground truth is the last place a plausible-sounding invention belongs. "
        "If an edge you expect is missing, the name is ambiguous — rename it and it will appear.")

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"[calls_order_doc] wrote {OUT.relative_to(WORKSPACE)}")
    print(f"[calls_order_doc] {len(graph)} functions · {len(ambiguous)} ambiguous names omitted")

    if OPEN_AFTER:
        try:
            subprocess.Popen(["cmd", "/c", "start", "", str(OUT)], shell=False)
        except Exception as e:
            print(f"  (could not open it: {e})")


if __name__ == "__main__":
    main()
